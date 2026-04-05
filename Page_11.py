import streamlit as st
from AdministrativeFunctions import change_page
from Category import category_codeID_label, select_categoryID, return_all_categories
from General_Functions import search_by_button, return_table, build_query, create_entry_user
from Ingredient import validate_ingredient_deep, get_ingredients_not_in_meal, return_all_ingredients
from Meal import collection_name, return_all_meals, meal_codeID_label, meal_createdAt_label, meal_name_label, make_meal, \
    full_entry_meal, update_meal, meal_notes_label, validate_meal, create_meal, delete_meal
from MealCombination import update_meal_combination, delete_meal_combination, create_meal_combination, \
    return_all_meal_combinations, meal_combination_quantity_label, validate_combination
from Menu import menu
from MongoDB_General_Functions import role_table
from Rule import rule_codeID_label, select_ruleID, return_all_rules
from UnitType import return_all_unit_types
from User import validate_user, user_collection_name, user_codeID_label
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Step 1: Initialize session state variables (first run only)
# Step 1a: Track the current page number
if "page" not in st.session_state:
    st.session_state.page = 1

# Step 1b: Track whether an error message should be shown
if "error_status" not in st.session_state:
    st.session_state.error_status = None

# Step 1c: Store the current signed-in user's CodeID
if "current_user" not in st.session_state:
    st.session_state.current_user = None

# Step 1d: Store the current error message (global handler reads this)
if "error" not in st.session_state:
    st.session_state.error = "You are doing great! Keep going."

# Step 1e: Store the CodeID to open in "full view"
if "open_code" not in st.session_state:
    st.session_state.open_code = None


def page_11_layout():
    # Step 2: Page entrypoint
    # Page 11: Meal search / management page.
    # Validates user, renders menu, and loads search UI + results.

    # Step 2a: Validate the currently signed-in user
    # Uses the CodeID stored in session state
    message, entry, status = validate_user(st.session_state.current_user)

    # Step 2b: If valid, render menu and enforce permissions
    if status:
        menu(entry)

        # Step 2c: Permission gate (admin only)
        if entry[0]["Role"] == role_table["Administrator"]:
            st.title(f"Collection {collection_name} Management")

            # Step 2d: Fetch entries and render search workflow
            all_entries = return_all_meals({})
            selected_entry = search(all_entries)

            # Step 2e: Alter, Add and Delete meals
            alter_meal(selected_entry, all_entries)
            add_meal()
            remove_meal(selected_entry, all_entries)

        else:
            # Step 2i: User valid, but not authorized
            st.session_state.error = f"[{user_collection_name}] Invalid Input: User does not have Administrator Privileges"
            st.session_state.error_status = False

    else:
        # Step 2f: Validation failed; store error for global handler
        st.session_state.error = message
        st.session_state.error_status = status


def search(entries):
    # Step 3: Search workflow (form + output)
    # Step 3a: Run search form and retrieve matching results
    outcome = search_form(entries)

    # Step 3b: Display search results
    display_results(outcome)

    # Step 3c: Return head result
    if len(outcome) >= 1:
        return outcome[0]
    else:
        return None


def search_form(entries):
    st.header("Search Form")

    # Step 3c: Build search UI layout
    with st.container(border=True):
        # Step 3d: Collect search filters from UI

        attribute = "CodeID"
        codeID = search_by_button(meal_codeID_label, return_table(attribute, entries, None, False),
                                  collection_name,
                                  attribute, "Search")

        attribute = "CreatedAt"
        createdAt = search_by_button(meal_createdAt_label, return_table(attribute, entries, None, True),
                                     collection_name, attribute, "Search")

        attribute = "UserID"
        userID = search_by_button(user_codeID_label, return_table(attribute, entries, None, False), collection_name,
                                  attribute, "Search")

        attribute = "CategoryID"
        categoryID = search_by_button(category_codeID_label, return_table(attribute, entries, None, False),
                                      collection_name,
                                      attribute, "Search")

        attribute = "RuleID"
        ruleID = search_by_button(rule_codeID_label, return_table(attribute, entries, None, False), collection_name,
                                  attribute, "Search")

        attribute = "Name"
        name = search_by_button(meal_name_label, return_table(attribute, entries, None, False), collection_name,
                                attribute, "Search")

        attribute = "Notes"
        notes = search_by_button(meal_notes_label, return_table(attribute, entries, None, False), collection_name,
                                 attribute, "Search")

        # Step 3e: Construct day-shaped object for query consistency
        query_like = make_meal(codeID, createdAt, userID, categoryID, name, notes, ruleID)

        # Step 3f: Show results
        show_result = st.checkbox("Show Results")

        # Step 3g: Build query and fetch matching results
        if show_result:
            query = build_query(query_like)
            return return_all_meals(query)
        else:
            return {}


def display_results(outcome):
    # Step 4: Results rendering
    st.header("Results")

    # Step 4a: Handle empty results
    if len(outcome) == 0:
        st.write("No results Found")
        return

    # Step 4b: Display count
    st.write(f"You have {len(outcome)} result(s)")

    # Step 4c: Render each result card
    pointer = 0
    while pointer < len(outcome):

        with st.container(border=True):
            full_entry_meal(outcome, pointer, True)

            # Step 4d: If first on list, marked as selected
            if pointer == 0:
                st.markdown(
                    """
                        <h1 style="text-align: center; font-size: 40px; font-weight: bold;"> Selected Entry </h1>
                    """,
                    unsafe_allow_html=True,
                )

        # Step 4e: Move pointer forward (prevents infinite loop)
        pointer += 1


def alter_meal_officially(codeID: str, userID: str, name: str, categoryID: str, notes: str, ruleID: str = None):
    # Step 1: Persist the update to the database (and capture status into session_state)
    st.session_state.error, _, st.session_state.error_status = update_meal(codeID, userID, name, categoryID,
                                                                           notes, ruleID)

    # Step 2: Renew Page to clean out error status if positive
    if st.session_state.error_status:
        change_page(st.session_state.page)


def alter_meal_combination_officially(codeID: str, userID: str, ingredientID: str, mealID: str, quantity: float):
    # Step 1: Persist the update to the database (and capture status into session_state)
    st.session_state.error, _, st.session_state.error_status = update_meal_combination(codeID, userID, ingredientID,
                                                                                       mealID,
                                                                                       quantity)

    # Step 2: Renew Page to clean out error status if positive
    if st.session_state.error_status:
        change_page(st.session_state.page)


def remove_meal_combination_officially(codeID: str, userID: str):
    # Step 1: Persist the update to the database (and capture status into session_state)
    st.session_state.error, _, st.session_state.error_status = delete_meal_combination(codeID, userID)

    # Step 2: Renew Page to clean out error status if positive
    if st.session_state.error_status:
        change_page(st.session_state.page)


def add_meal_combination_officially(userID: str, ingredientID: str, mealID: str, quantity: float):
    # Step 1: Persist the update to the database (and capture status into session_state)
    st.session_state.error, _, st.session_state.error_status = create_meal_combination(userID, ingredientID,
                                                                                       mealID,
                                                                                       quantity)

    # Step 2: Renew Page to clean out error status if positive
    if st.session_state.error_status:
        change_page(st.session_state.page)


def alter_meal(entry, entries):
    # Step 1: Set page header based on whether we’re editing a known entry or selecting one
    if entry is None:
        st.header(f"Alter {collection_name} Item")
        name_set, notes_set = "", ""
    else:
        st.header(f"Alter Item {entry['CodeID']}")
        name_set, notes_set = entry['Name'], entry['Notes']

    # Step 2: Wrap the form in a bordered container for UI grouping
    with st.container(border=True):

        # Step 3: If no entry was provided, ask the user which CodeID they want to alter
        if entry is None:
            attribute = "CodeID"
            codeID = search_by_button(meal_codeID_label, return_table(attribute, entries, [], False),
                                      collection_name, attribute, "Alter")
        # Step 4: Otherwise, CodeID is already known from entry
        else:

            codeID = st.text_input(meal_codeID_label, value=entry['CodeID'], key="alter_meal_codeID",
                                   disabled=True)

        # Step 5: Alter the Category for this meal (pre-filled if entry exists, depending on alter_ruleID)
        categoryID = select_categoryID(entry, 1)

        # Step 6: Pick/alter the rule for this meal (pre-filled if entry exists, depending on alter_ruleID)
        ruleID = select_ruleID(entry, 1)

        # Step 7: Input for meal name (pre-filled if entry exists)
        name = st.text_input(meal_name_label, value=name_set, key="alter_meal_name")

        # Step 8: Input for meal name (pre-filled if entry exists)
        notes = st.text_input(meal_notes_label, value=notes_set, key="alter_meal_notes")

        # Step 9: Validate meal exists
        meal_message, meal_entry, meal_status = validate_meal(codeID)

        if meal_status:

            # Step 10: Show and alter ingredients
            st.subheader(f"Ingredients for {name_set}")

            # Step 11: Gather existing ingredients
            existing_ingredients = return_all_meal_combinations({'MealID': codeID})

            for combination in existing_ingredients:

                # Step 12: Validate ingredient existence through combination
                status, message, unit_type, ingredient = validate_combination(combination)

                if status:

                    # Step 13: Show existing ingredient to be altered or deleted
                    show_ingredient(combination, ingredient, unit_type, codeID, "Exists")

                else:

                    st.write(message)

            # Step 14: Gather new ingredients
            other_ingredients = get_ingredients_not_in_meal(existing_ingredients)

            for other_ingredient in other_ingredients:

                # Step 15: Validate ingredient existence
                status, message, unit_type, ingredient = validate_ingredient_deep(other_ingredient)

                if status:

                    # Step 16: Show ingredient to be added
                    show_ingredient(None, ingredient, unit_type, codeID, "New")

                else:

                    st.write(message)

        # Step 18: Submit button - triggers DB update via callback, passing the collected inputs
        st.button(
            f"Update Entry {codeID}",
            use_container_width=True,
            on_click=alter_meal_officially,
            args=[codeID, st.session_state.current_user, name, categoryID, notes, ruleID],
            key=f"alter_meal"
        )


def show_ingredient(entry, ingredient, unit_type: str, meal: str, action: str):
    if action == "Exists":
        codeID = entry['CodeID']
        text_input_column, text_column, update_column, delete_column = st.columns(4,
                                                                                  vertical_alignment="center")
        quantity_set = float(entry['Quantity'])
    else:
        codeID = ingredient['CodeID']
        text_input_column, text_column, add_column = st.columns(3, vertical_alignment="center")
        quantity_set = 0.00

    with text_input_column:
        quantity = st.number_input(
            meal_combination_quantity_label,
            label_visibility="collapsed",
            value=quantity_set,
            min_value=0.00,
            step=0.05,
            key=f"{meal}_{action}_{codeID}_quantity",
        )

    with text_column:
        st.write(f"{unit_type}(s) of {ingredient['Name']}")

    if action == "Exists":
        with update_column:
            st.button(
                f"Update {codeID}",
                use_container_width=True,
                on_click=alter_meal_combination_officially,
                args=[codeID, st.session_state.current_user, entry['IngredientID'],
                      entry['MealID'], quantity],
                key=f"update_{meal}_{action}_{codeID}"
            )
        with delete_column:
            st.button(
                f"Delete {codeID}",
                use_container_width=True,
                on_click=remove_meal_combination_officially,
                args=[codeID, st.session_state.current_user],
                key=f"delete_{meal}_{action}_{codeID}"
            )

    else:

        with add_column:
            st.button(
                f"Add {codeID}",
                use_container_width=True,
                on_click=add_meal_combination_officially,
                args=[st.session_state.current_user, ingredient['CodeID'], meal, quantity],
                key=f"add_{meal}_{action}_{codeID}"
            )


def add_meal_officially(name: str, userID: str, categoryID: str, notes: str, ruleID: str = None):
    # Step 1: Persist the create operation to the database (and capture status into session_state)
    st.session_state.error, _, st.session_state.error_status = create_meal(name, userID, categoryID, notes, ruleID)

    # Step 2: Renew Page to clean out error status if positive
    if st.session_state.error_status:
        change_page(st.session_state.page)


def add_meal():
    # Step 1: Page header for add flow
    st.header(f"Add {collection_name} Item")

    # Step 2: Wrap the form in a bordered container for UI grouping
    with st.container(border=True):
        # Step 3: Show userID (current user) but disable editing
        userID = create_entry_user(user_codeID_label, collection_name, st.session_state.current_user)

        # Step 4: Alter the Category for this meal (pre-filled if entry exists, depending on alter_ruleID)
        categoryID = select_categoryID(None, 2)

        # Step 5: Pick rule for this new ingredient
        ruleID = select_ruleID(None, 2)

        # Step 6: Input for meal name (pre-filled if entry exists)
        name = st.text_input(meal_name_label, key="add_meal_name")

        # Step 7: Input for meal name (pre-filled if entry exists)
        notes = st.text_input(meal_notes_label, key="add_meal_notes")

        # Step 8: Submit button - should create a new entry
        st.button(
            f"Add Entry",
            use_container_width=True,
            on_click=add_meal_officially,
            args=[name, userID, categoryID, notes, ruleID],
            key=f"add_meal"
        )


def remove_meal_officially(codeID: str, userID: str):
    # Step 1: Persist the delete operation to the database (and capture status into session_state)
    st.session_state.error, _, st.session_state.error_status = delete_meal(codeID, userID)

    # Step 2: Renew Page to clean out error status if positive
    if st.session_state.error_status:
        change_page(st.session_state.page)


def remove_meal(entry, entries):
    # Step 1: Set page header based on whether we’re deleting a known entry or selecting one
    if entry is None:
        st.header(f"Delete {collection_name} Item")
    else:
        st.header(f"Delete Item {entry['CodeID']}")

    # Step 2: Wrap the delete flow in a bordered container for UI grouping
    with st.container(border=True):

        # Step 3: If no entry was provided, ask the user which CodeID they want to delete
        if entry is None:
            attribute = "CodeID"
            codeID = search_by_button(meal_codeID_label, return_table(attribute, entries, [], False),
                                      collection_name, attribute, "Delete")

        # Step 4: Otherwise show CodeID as read-only (already known)
        else:
            codeID = st.text_input(
                meal_codeID_label,
                value=entry['CodeID'],
                key="delete_ingredient_codeID",
                disabled=True
            )

        # Step 5: Validate that this ingredient exists / can be deleted, and fetch full entry data
        ingredient_message, ingredient_entry, ingredient_status = validate_meal(codeID)

        # Step 6: If valid, show full entry preview and present delete button
        if ingredient_status:
            full_entry_meal(ingredient_entry, 0, False)

            # Step 7: Confirm delete action with a button; calls delete callback with CodeID and current user
            st.button(
                f"Delete Entry {codeID}",
                use_container_width=True,
                on_click=remove_meal_officially,
                args=[codeID, st.session_state.current_user],
                key=f"delete_meal"
            )


def manage_meals():
    # Step 5a: Retrieve all meals belonging to the current user
    entries = return_all_meals({"UserID": st.session_state.current_user})

    # Step 5b: Present the search/action interface for meal management
    option = search_by_button(meal_name_label,
                              return_table("Name", entries,
                                           ["Add", "Show All"], False), collection_name, "Name",
                              "Search")

    if option == "Show All":

        # Step 5d: Make and download cookbook word file
        cookbook = make_cookbook(entries)
        cookbook = sorted(cookbook, key=lambda x: (x.get("Title") or "").lower())
        docx_file = build_cookbook_docx(cookbook)
        st.download_button(
            label="Download Cookbook",
            data=docx_file,
            use_container_width=True,
            file_name="my_cookbook.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Step 5c: Display all meals if the user selects "Show All"
        display_results(entries)

    # Step 5e: Start the meal creation workflow if the user selects "Add"
    elif option == "Add":
        add_meal()

    else:
        # Step 5f: Retrieve the selected meal and allow modification or removal
        meal_entry = return_all_meals({"UserID": st.session_state.current_user, "Name": option})

        if len(meal_entry) >= 1:

            # Step 5g: Make and download cookbook word file
            cookbook = make_cookbook(meal_entry)
            docx_file = build_cookbook_docx(cookbook)
            st.download_button(
                label=f"Download {meal_entry[0]['Name']}",
                data=docx_file,
                use_container_width=True,
                file_name="my_cookbook.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Step 5h: Ensure the selected meal exists before altering or removing it
            alter_meal(meal_entry[0], meal_entry)
            remove_meal(meal_entry[0], meal_entry)


def make_cookbook(entries):
    # Step 1: Validate user once and get username for "CreatedBy"
    message, user, status = validate_user(st.session_state.current_user)
    if not status:
        return []

    created_by = user[0]["Username"]

    # Step 2: Extract all MealIDs from incoming entries
    # These are used to fetch all related data in bulk (instead of per-entry calls)
    meal_ids = [entry["CodeID"] for entry in entries if entry.get("CodeID")]
    if not meal_ids:
        return []

    # Step 3: Fetch all meals in one query and build lookup
    meals = return_all_meals({
        "CodeID": {"$in": meal_ids},
        "UserID": st.session_state.current_user
    })
    meal_lookup = {meal["CodeID"]: meal for meal in meals}

    # Step 4: Collect all CategoryIDs and RuleIDs from meals
    # These will be used for batch fetching
    category_ids = list({
        meal["CategoryID"]
        for meal in meals
        if meal.get("CategoryID") is not None
    })

    rule_ids = list({
        meal["RuleID"]
        for meal in meals
        if meal.get("RuleID") is not None
    })

    # Step 5: Fetch categories and build lookup
    categories = return_all_categories({
        "CodeID": {"$in": category_ids}
    }) if category_ids else []
    category_lookup = {category["CodeID"]: category for category in categories}

    # Step 6: Fetch rules and build lookup
    rules = return_all_rules({
        "CodeID": {"$in": rule_ids}
    }) if rule_ids else []
    rule_lookup = {rule["CodeID"]: rule for rule in rules}

    # Step 7: Fetch all meal combinations (ingredient links) in one query
    meal_combinations = return_all_meal_combinations({
        "MealID": {"$in": meal_ids},
        "UserID": st.session_state.current_user
    })

    # Step 8: Organize combinations by MealID and collect all IngredientIDs
    combinations_by_meal_id = {}
    ingredient_ids = set()

    for combo in meal_combinations:
        meal_id = combo.get("MealID")
        ingredient_id = combo.get("IngredientID")

        if meal_id is not None:
            combinations_by_meal_id.setdefault(meal_id, []).append(combo)

        if ingredient_id is not None:
            ingredient_ids.add(ingredient_id)

    # Step 9: Fetch all ingredients in one query and build lookup
    ingredients = return_all_ingredients({
        "CodeID": {"$in": list(ingredient_ids)}
    }) if ingredient_ids else []
    ingredient_lookup = {ingredient["CodeID"]: ingredient for ingredient in ingredients}

    # Step 10: Collect UnitTypeIDs from ingredients and fetch them
    unit_type_ids = list({
        ingredient["UnitTypeID"]
        for ingredient in ingredients
        if ingredient.get("UnitTypeID") is not None
    })

    units = return_all_unit_types({
        "CodeID": {"$in": unit_type_ids}
    }) if unit_type_ids else []
    unit_lookup = {unit["CodeID"]: unit for unit in units}

    # Step 11: Build the final cookbook structure
    cookbook = []

    for entry in entries:
        meal_id = entry.get("CodeID")
        meal = meal_lookup.get(meal_id)

        # Step 11.1: Skip if meal is missing
        if not meal:
            continue

        # Step 11.2: Resolve category and rule (if they exist)
        category = category_lookup.get(meal.get("CategoryID"))
        rule = rule_lookup.get(meal.get("RuleID"))

        # Step 11.3: Format rule text (e.g. "2 per week")
        rule_text = None
        if rule:
            quantity = rule.get("Quantity")
            per = rule.get("Per")
            if quantity is not None and per:
                rule_text = f"{quantity} per {per}"

        # Step 11.4: Build ingredient list for this meal
        ingredients_list = []

        for combo in combinations_by_meal_id.get(meal_id, []):
            ingredient = ingredient_lookup.get(combo.get("IngredientID"))
            if not ingredient:
                continue

            unit = unit_lookup.get(ingredient.get("UnitTypeID"))
            if not unit:
                continue

            ingredients_list.append({
                "Ingredient": ingredient.get("Name"),
                "Quantity": combo.get("Quantity"),
                "UnitType": unit.get("Name")
            })

        # Step 11.5: Assemble final cookbook row
        cookbook.append({
            "Title": meal.get("Name"),
            "Category": category.get("Name") if category else None,
            "Created": meal.get("CreatedAt"),
            "CreatedBy": created_by,
            "Rule": rule_text,
            "Ingredients": ingredients_list,
            "Notes": meal.get("Notes")
        })

    # Step 12: Return final cookbook and error counter
    return cookbook


def build_cookbook_docx(cookbook):
    # Step 1: Create document
    document = Document()

    # Step 2: Add title page
    title = document.add_heading("My Cookbook", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")  # spacing

    # Step 3: Loop through recipes
    for index, meal in enumerate(cookbook):
        # Step 3.1: Meal title
        document.add_heading(meal.get("Title", "Unnamed Meal"), level=1)

        # Step 3.2: Metadata (Category, Created, Rule)
        if meal.get("Category"):
            document.add_paragraph(f"Category: {meal['Category']}")

        if meal.get("Created"):
            document.add_paragraph(f"Created: {meal['Created']}")

        if meal.get("CreatedBy"):
            document.add_paragraph(f"By: {meal['CreatedBy']}")

        if meal.get("Rule"):
            document.add_paragraph(f"Rule: {meal['Rule']}")

        # Step 3.3: Ingredients section
        document.add_heading("Ingredients", level=2)

        ingredients = meal.get("Ingredients", [])
        if ingredients:
            for ing in ingredients:
                text = f"{ing.get('Quantity', '')} {ing.get('UnitType', '')} {ing.get('Ingredient', '')}"
                document.add_paragraph(text.strip(), style="List Bullet")
        else:
            document.add_paragraph("—")

        # Step 3.4: Notes section
        document.add_heading("Notes", level=2)

        notes = meal.get("Notes")
        document.add_paragraph(notes if notes else "—")

        # Step 3.5: Page break between meals (not after last)
        if index < len(cookbook) - 1:
            document.add_page_break()

    # Step 4: Save to memory (for Streamlit download)
    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return output
